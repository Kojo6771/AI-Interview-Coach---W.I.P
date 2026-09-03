from io import BytesIO
from docx import Document
from pypdf import PdfReader

#extract text from PDF
def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))

    text_parts = []

    for page in reader.pages:
        text = page.extract_text()

        if text:
            text_parts.append(text)

    return "\n".join(text_parts).strip()

#extract text from DOCX
def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))

    text_parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            text_parts.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    text_parts.append(text)

    return "\n".join(text_parts).strip()

#extract text from CV based on file extension
def extract_cv_text(
    file_bytes: bytes,
    file_extension: str
) -> str:

    if file_extension == ".pdf":
        return extract_pdf_text(file_bytes)

    if file_extension == ".docx":
        return extract_docx_text(file_bytes)

    raise ValueError(
        f"Unsupported file type: {file_extension}"
    )